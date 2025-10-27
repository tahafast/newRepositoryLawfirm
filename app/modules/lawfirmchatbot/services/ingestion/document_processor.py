"""Robust document processing service for PDF, DOCX, DOC, and TXT files."""

from typing import List, Optional
from app.modules.lawfirmchatbot.services._lc_compat import ensure_Document, get_recursive_splitter
from app.modules.lawfirmchatbot.services.text_cleaner import clean_and_validate_text
import io
import math
import logging
import os
import re
import unicodedata
from pathlib import Path

logger = logging.getLogger(__name__)

Document = ensure_Document()


def _clean_unicode_text(text: str) -> str:
    """Clean problematic Unicode characters that cause JSON serialization issues."""
    if not text:
        return text
    
    try:
        # Normalize Unicode characters
        text = unicodedata.normalize('NFKD', text)
        
        # Remove surrogate pairs and problematic characters
        text = re.sub(r'[\ud800-\udfff]', '', text)  # Remove surrogate pairs
        text = re.sub(r'[\ufffe\uffff]', '', text)   # Remove non-characters
        
        # Replace mathematical symbols that cause issues
        text = re.sub(r'[\U0001d400-\U0001d7ff]', '[MATH]', text)  # Mathematical symbols
        text = re.sub(r'[\U00010000-\U0010ffff]', '[SYMBOL]', text)  # High Unicode planes
        
        # Ensure the text can be encoded to UTF-8
        text = text.encode('utf-8', errors='ignore').decode('utf-8')
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    except Exception as e:
        logger.warning(f"Unicode cleaning failed: {e}")
        # Fallback: encode/decode to remove problematic characters
        return text.encode('ascii', errors='ignore').decode('ascii')


def process_document(file_path: str, filename: str, chunk_size: int = 1000) -> List[Document]:
    """Process document with robust error handling - functional interface."""
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Extract text based on file type
        text = _extract_text_safe(file_path)
        if not text or not text.strip():
            raise ValueError(f"No text content extracted from {filename}")

        # Clean problematic Unicode characters
        text = _clean_unicode_text(text)

        # Run environment-aware cleaning before chunking
        cleaned_text, clean_meta = clean_and_validate_text(text, filename=filename)
        if not clean_meta.get("kept", False):
            logger.warning(
                "[document-processor] Skipping embedding for %s (quality=%s, alpha=%.2f, removed=%.2f)",
                filename,
                clean_meta.get("quality"),
                clean_meta.get("alpha_ratio", 0.0),
                clean_meta.get("removed_ratio", 0.0),
            )
            return []
        text = cleaned_text
        
        # Get page count
        total_pages = _get_page_count_safe(file_path)
        
        # Create text splitter with optimized chunking (900-1200 tokens ≈ 3600-4800 chars)
        # Using 4000 chars as middle ground with 480 overlap (~120 tokens)
        text_splitter = get_recursive_splitter(
            chunk_size=4000,  # ~1000 tokens for better context
            chunk_overlap=480,  # ~120 tokens overlap
            length_function=len,
            add_start_index=True,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Create chunks
        chunks = text_splitter.create_documents(
            texts=[text],
            metadatas=[{
                "source": filename,
                "total_pages": total_pages,
                "page": 1,
                "file_type": Path(filename).suffix.lower(),
                "clean_version": clean_meta.get("clean_version"),
                "clean_quality": clean_meta.get("quality"),
                "clean_removed_ratio": clean_meta.get("removed_ratio"),
                "clean_alpha_ratio": clean_meta.get("alpha_ratio"),
            }]
        )

        for chunk in chunks:
            chunk.page_content = chunk.page_content.strip()
        
        # Distribute chunks across pages
        _assign_page_numbers(chunks, text, total_pages)
        
        logger.info(f"Successfully processed {filename}: {len(chunks)} chunks across {total_pages} pages")
        return chunks
        
    except Exception as e:
        logger.error(f"Error processing file {filename}: {str(e)}", exc_info=True)
        raise


def _extract_text_safe(file_path: str) -> str:
    """Extract text with comprehensive error handling."""
    file_ext = Path(file_path).suffix.lower()
    
    try:
        if file_ext == '.pdf':
            return _extract_from_pdf_safe(file_path)
        elif file_ext == '.docx':
            return _extract_from_docx_safe(file_path)
        elif file_ext == '.doc':
            return _extract_from_doc_safe(file_path)
        elif file_ext == '.txt':
            return _extract_from_txt_safe(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    except Exception as e:
        logger.error(f"Text extraction failed for {file_path}: {str(e)}")
        raise


def _extract_from_pdf_safe(file_path: str) -> str:
    """Extract text from PDF with multiple fallback methods."""
    text = ""
    
    # Try PyPDF2 first
    try:
        import PyPDF2
        with open(file_path, 'rb') as file:
            pdf = PyPDF2.PdfReader(file)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    # Clean Unicode characters immediately
                    page_text = _clean_unicode_text(page_text)
                    text += page_text + '\n'
        
        if text.strip():
            return text
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {e}")
    
    if not text.strip():
        raise ValueError("Could not extract text from PDF using available methods")
    
    return text


def _extract_from_docx_safe(file_path: str) -> str:
    """Extract text from DOCX with error handling."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        
        if not text.strip():
            raise ValueError("No text content found in DOCX file")
        
        # Clean Unicode characters
        text = _clean_unicode_text(text)
        return text
    except ImportError:
        raise ImportError("python-docx not installed. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise


def _extract_from_doc_safe(file_path: str) -> str:
    """Extract text from legacy DOC files using best-effort strategies."""
    try:
        with open(file_path, 'rb') as file:
            raw_bytes = file.read()
    except Exception as e:
        logger.error(f"DOC extraction failed: {e}")
        raise

    # Try textract first (covers legacy encodings / embedded fonts)
    try:
        import textract  # type: ignore

        textract_bytes = textract.process(io.BytesIO(raw_bytes), extension="doc")
        if isinstance(textract_bytes, bytes):
            textract_text = textract_bytes.decode("utf-8", "ignore")
        else:
            textract_text = str(textract_bytes or "")
        textract_text = _clean_doc_binary_text(textract_text)
        if textract_text.strip():
            return textract_text
    except ImportError:
        pass
    except Exception as tex_err:
        logger.warning(f"textract DOC extraction failed: {tex_err}")

    stripped = raw_bytes.lstrip()

    # DOC files saved as RTF are common; detect and convert
    if stripped.startswith(b'{\\rtf'):
        rtf_text = _rtf_bytes_to_text(raw_bytes)
        if rtf_text and rtf_text.strip():
            return rtf_text
        logger.warning("RTF conversion for DOC file returned empty text; falling back to binary decode")

    # Fallback: decode binary content and strip control characters
    for encoding in ('cp1252', 'utf-8', 'latin-1'):
        try:
            decoded = raw_bytes.decode(encoding)
            cleaned = _clean_doc_binary_text(decoded)
            if cleaned.strip():
                return cleaned
        except UnicodeDecodeError:
            continue

    # If nothing produced usable text, raise error
    raise ValueError("Could not extract text content from DOC file")


def _extract_from_txt_safe(file_path: str) -> str:
    """Extract text from TXT with encoding detection."""
    encodings = ['utf-8', 'utf-16', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
                if text.strip():
                    # Clean Unicode characters
                    text = _clean_unicode_text(text)
                    return text
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception as e:
            logger.warning(f"Failed to read with {encoding}: {e}")
            continue
    
    raise ValueError("Could not decode text file with any supported encoding")


def _get_page_count_safe(file_path: str) -> int:
    """Get page count with fallback methods."""
    file_ext = Path(file_path).suffix.lower()
    
    try:
        if file_ext == '.pdf':
            return _get_pdf_page_count_safe(file_path)
        elif file_ext == '.docx':
            return _get_docx_page_count_safe(file_path)
        elif file_ext == '.doc':
            return _get_doc_page_count_safe(file_path)
        elif file_ext == '.txt':
            return _get_txt_page_count_safe(file_path)
        else:
            return 1
    except Exception as e:
        logger.warning(f"Page count estimation failed: {e}")
        return 1


def _get_pdf_page_count_safe(file_path: str) -> int:
    """Get PDF page count with fallbacks."""
    try:
        import PyPDF2
        with open(file_path, 'rb') as file:
            pdf = PyPDF2.PdfReader(file)
            return len(pdf.pages)
    except Exception:
        # Fallback: estimate based on file size
        file_size = os.path.getsize(file_path)
        return max(1, file_size // 50000)  # Rough estimate


def _get_docx_page_count_safe(file_path: str) -> int:
    """Get DOCX page count estimation."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        total_chars = sum(len(paragraph.text) for paragraph in doc.paragraphs)
        return max(1, math.ceil(total_chars / 3000))
    except Exception:
        return 1


def _get_txt_page_count_safe(file_path: str) -> int:
    """Get TXT page count estimation."""
    try:
        file_size = os.path.getsize(file_path)
        return max(1, math.ceil(file_size / 3000))
    except Exception:
        return 1


def _assign_page_numbers(chunks: List[Document], text: str, total_pages: int) -> None:
    """Assign page numbers to chunks."""
    if total_pages <= 1:
        return
    
    total_chars = len(text)
    chars_per_page = total_chars / total_pages if total_pages > 0 else total_chars
    
    for chunk in chunks:
        char_position = chunk.metadata.get('start_index', 0)
        page_num = min(total_pages, max(1, math.ceil((char_position + 1) / chars_per_page)))
        chunk.metadata["page"] = page_num


def _rtf_bytes_to_text(raw_bytes: bytes) -> str:
    """Convert RTF bytes to plain text."""
    for encoding in ("utf-8", "utf-16", "utf-16le", "utf-16be", "cp1252", "latin-1"):
        try:
            candidate = raw_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        candidate = raw_bytes.decode("utf-8", errors="ignore")

    try:
        from striprtf.striprtf import rtf_to_text
        return rtf_to_text(candidate)
    except Exception as e:
        logger.warning(f"striprtf conversion failed ({e}); using regex fallback for RTF text extraction")
        text = re.sub(r'\\par[d]?', '\n', candidate)
        text = re.sub(r'\\[A-Za-z]+\d*(?:\s|)', ' ', text)
        text = text.replace('{', ' ').replace('}', ' ')
        return text


def _clean_doc_binary_text(text: str) -> str:
    """Remove binary control characters and normalize whitespace from DOC content."""
    if not text:
        return text

    text = text.replace('\x00', '')
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = ''.join(ch if ch.isprintable() or ch in '\n\t' else ' ' for ch in text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    return text.strip()


def _get_doc_page_count_safe(file_path: str) -> int:
    """Estimate DOC page count using file size heuristic."""
    try:
        file_size = os.path.getsize(file_path)
        # DOC files tend to be denser; use slightly larger divisor than TXT
        return max(1, math.ceil(file_size / 4500))
    except Exception:
        return 1
